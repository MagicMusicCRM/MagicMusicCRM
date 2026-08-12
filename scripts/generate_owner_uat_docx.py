"""Generate the owner mega-UAT pre-final signing document.

The source of truth remains the Markdown result matrix. This script turns its
100 rows into a deterministic Word form without changing scenario statuses.
"""

from __future__ import annotations

import argparse
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
TABLE_COLUMNS_DXA = (1008, 4320, 1008, 1800, 1224)

INK = "172033"
MUTED = "667085"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
GREEN = "E7F4EA"
GOLD = "FFF4D6"
RED = "FDE7E7"


def set_run_font(
    run,
    *,
    size: float,
    color: str = INK,
    bold: bool = False,
    italic: bool = False,
) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def configure_styles(document: Document) -> None:
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(INK)
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.line_spacing = 1.0

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = RGBColor.from_string(MUTED)
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(18)
    subtitle.paragraph_format.line_spacing = 1.15

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_page(document: Document) -> None:
    for section in document.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run("Страница ")
    set_run_font(run, size=9, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    field_run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), MUTED)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "18")
    properties.extend([color, size])
    field_run.append(properties)
    text = OxmlElement("w:t")
    text.text = "1"
    field_run.append(text)
    field.append(field_run)
    paragraph._p.append(field)


def configure_header_footer(document: Document) -> None:
    for section in document.sections:
        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run("MagicMusicCRM · OWNER MEGA-UAT · PRE-FINAL")
        set_run_font(run, size=8.5, color=MUTED, bold=True)

        footer = section.footer
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_paragraph.paragraph_format.space_before = Pt(0)
        footer_paragraph.paragraph_format.space_after = Pt(0)
        add_page_field(footer_paragraph)


def add_numbering(document: Document) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    number_format = OxmlElement("w:numFmt")
    number_format.set(qn("w:val"), "bullet")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•")
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    paragraph_properties = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    paragraph_properties.extend([tabs, indent, spacing])
    level.extend([start, number_format, level_text, justification, paragraph_properties])
    abstract.append(level)
    numbering.append(abstract)

    instance = OxmlElement("w:num")
    instance.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    instance.append(abstract_ref)
    numbering.append(instance)
    return num_id


def add_bullet(document: Document, text: str, num_id: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    properties = paragraph._p.get_or_add_pPr()
    num_properties = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    num_properties.extend([level, number])
    properties.append(num_properties)
    run = paragraph.add_run(text)
    set_run_font(run, size=11)


def add_labeled_paragraph(document: Document, label: str, value: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run, size=10.5, bold=True)
    value_run = paragraph.add_run(value)
    set_run_font(value_run, size=10.5)


def shade_paragraph(paragraph, fill: str) -> None:
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "20")
    left.set(qn("w:space"), "10")
    left.set(qn("w:color"), BLUE)
    borders.append(left)
    properties.append(borders)


def set_cell_margins(cell, *, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for side, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)


def set_table_geometry(table, widths_dxa: tuple[int, ...]) -> None:
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError("Table widths must sum to 9360 DXA")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    properties = table._tbl.tblPr
    width = properties.first_child_found_in("w:tblW")
    width.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    width.set(qn("w:type"), "dxa")
    layout = properties.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")
    indent = properties.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    indent.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(value))
        grid.append(column)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[index] / 1440)
            cell_properties = cell._tc.get_or_add_tcPr()
            cell_width = cell_properties.first_child_found_in("w:tcW")
            cell_width.set(qn("w:w"), str(widths_dxa[index]))
            cell_width.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def repeat_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    properties.append(marker)


def prevent_row_split(row) -> None:
    properties = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:cantSplit")
    marker.set(qn("w:val"), "true")
    properties.append(marker)


def format_cell(cell, *, size: float = 8.5, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        for run in paragraph.runs:
            set_run_font(run, size=size, bold=bold)


def parse_matrix(path: Path) -> list[tuple[str, str, str, str]]:
    rows: list[tuple[str, str, str, str]] = []
    for line in path.read_text(encoding="utf-8").splitlines():
        if not line.startswith("| UAT-"):
            continue
        cells = [value.strip() for value in line.split("|")[1:-1]]
        if len(cells) != 4:
            raise ValueError(f"Unexpected UAT row: {line}")
        rows.append((cells[0], cells[1], cells[2], cells[3]))
    if len(rows) != 100:
        raise ValueError(f"Expected 100 UAT rows, found {len(rows)}")
    if len({row[0] for row in rows}) != 100:
        raise ValueError("UAT identifiers must be unique")
    return rows


def add_summary_table(document: Document, counts: Counter[str]) -> None:
    statuses = ("PASS", "PARTIAL", "PENDING", "FAIL", "BLOCKED")
    table = document.add_table(rows=2, cols=5)
    table.style = "Table Grid"
    widths = (1872, 1872, 1872, 1872, 1872)
    for index, status in enumerate(statuses):
        table.cell(0, index).text = status
        table.cell(1, index).text = str(counts.get(status, 0))
        shade_cell(table.cell(0, index), LIGHT_BLUE)
        format_cell(table.cell(0, index), size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        format_cell(table.cell(1, index), size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_table_geometry(table, widths)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)


def add_decision_matrix(document: Document, rows: list[tuple[str, str, str, str]]) -> None:
    fills = {
        "PASS": GREEN,
        "PARTIAL": GOLD,
        "PENDING": LIGHT_GRAY,
        "FAIL": RED,
        "BLOCKED": RED,
    }
    # Keep each signing table deliberately below a full page. LibreOffice can
    # otherwise place the following manual break inside the page header area.
    chunks = [rows[index : index + 15] for index in range(0, len(rows), 15)]
    for chunk_index, chunk in enumerate(chunks):
        heading = document.add_paragraph(
            f"Матрица решений владельца · {chunk_index + 1}/{len(chunks)}",
            style="Heading 1",
        )
        heading.paragraph_format.page_break_before = True
        intro = document.add_paragraph(
            "Отметьте один итог в четвёртой колонке и заполните инициалы/дату. "
            "Подробные evidence-ссылки остаются в исходной Markdown-матрице."
        )
        intro.paragraph_format.space_after = Pt(10)

        table = document.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        headers = ("ID", "Сценарий", "Тех. статус", "Решение владельца", "Инициалы / дата")
        for index, value in enumerate(headers):
            table.cell(0, index).text = value
            shade_cell(table.cell(0, index), LIGHT_BLUE)
            format_cell(table.cell(0, index), size=8.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        repeat_header(table.rows[0])

        for uat_id, scenario, status, _evidence in chunk:
            row = table.add_row()
            prevent_row_split(row)
            values = (uat_id, scenario, status, "PASS / FAIL / BLOCKED", "________\n________")
            for index, value in enumerate(values):
                row.cells[index].text = value
            shade_cell(row.cells[2], fills.get(status, LIGHT_GRAY))
            format_cell(row.cells[0], size=8.3, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            format_cell(row.cells[1], size=8.5)
            format_cell(row.cells[2], size=8.0, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            format_cell(row.cells[3], size=7.7, align=WD_ALIGN_PARAGRAPH.CENTER)
            format_cell(row.cells[4], size=8.0, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_table_geometry(table, TABLE_COLUMNS_DXA)


def build_document(source: Path, output: Path) -> None:
    rows = parse_matrix(source)
    counts = Counter(row[2] for row in rows)
    document = Document()
    configure_styles(document)
    configure_page(document)
    configure_header_footer(document)
    bullet_num_id = add_numbering(document)

    properties = document.core_properties
    properties.title = "MagicMusicCRM — owner mega-UAT pre-final"
    properties.subject = "Матрица итоговой production-приёмки"
    properties.author = "MagicMusicCRM"
    properties.keywords = "UAT, owner acceptance, MagicMusicCRM"
    properties.comments = "Generated from v7-owner-production-mega-uat-result.md"

    kicker = document.add_paragraph()
    kicker.paragraph_format.space_before = Pt(12)
    kicker.paragraph_format.space_after = Pt(4)
    kicker_run = kicker.add_run("PRE-FINAL · OWNER REVIEW")
    set_run_font(kicker_run, size=9.5, color=BLUE, bold=True)

    document.add_paragraph("Итоговая матрица production-приёмки", style="Title")
    document.add_paragraph(
        "MagicMusicCRM 1.5.1+181 · 100 утверждённых сценариев",
        style="Subtitle",
    )
    add_labeled_paragraph(document, "Дата снимка", "12 августа 2026")
    add_labeled_paragraph(document, "Production-кандидат", "Client 1.5.1+181 · Server b04f177")
    add_labeled_paragraph(document, "Источник", "docs/audits/v7-owner-production-mega-uat-result.md")
    add_labeled_paragraph(document, "Назначение", "форма покомпонентной приёмки владельцем")

    callout = document.add_paragraph()
    callout.paragraph_format.left_indent = Inches(0.14)
    callout.paragraph_format.right_indent = Inches(0.08)
    callout.paragraph_format.space_before = Pt(12)
    callout.paragraph_format.space_after = Pt(14)
    shade_paragraph(callout, CALLOUT)
    callout_run = callout.add_run(
        "Документ не является финальной подписью. Технических PENDING нет, но "
        "90 строк остаются PARTIAL до production owner-UAT. Пустое решение "
        "владельца не считается PASS."
    )
    set_run_font(callout_run, size=10.5, bold=True)

    document.add_paragraph("Текущий технический снимок", style="Heading 1")
    add_summary_table(document, counts)

    document.add_paragraph("Как заполнять", style="Heading 1")
    add_bullet(document, "Для каждой строки отметить ровно одно решение: PASS, FAIL или BLOCKED.", bullet_num_id)
    add_bullet(document, "Поставить инициалы и дату напротив каждой проверенной строки.", bullet_num_id)
    add_bullet(document, "Для FAIL/BLOCKED приложить номер дефекта и краткое воспроизведение.", bullet_num_id)
    add_bullet(document, "Подписывать общий итог только после заполнения всех 100 строк.", bullet_num_id)

    document.add_paragraph("Границы документа", style="Heading 1")
    boundary = document.add_paragraph(
        "Технический статус отражает автоматизированные и локальные доказательства на дату снимка. "
        "Решение владельца фиксирует фактическое поведение production Release с реальными ролями и данными."
    )
    boundary.paragraph_format.space_after = Pt(4)

    add_decision_matrix(document, rows)

    signature_heading = document.add_paragraph("Итоговая подпись", style="Heading 1")
    signature_heading.paragraph_format.page_break_before = True
    document.add_paragraph(
        "Заполняется владельцем после завершения всех строк и сверки приложенных доказательств."
    )
    for label in (
        "Итог PASS",
        "Итог FAIL",
        "Итог BLOCKED",
        "Открытые дефекты / ссылки",
        "Решение о выпуске",
    ):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(6)
        label_run = paragraph.add_run(f"{label}: ")
        set_run_font(label_run, size=11, bold=True)
        line_run = paragraph.add_run("____________________________________________________________")
        set_run_font(line_run, size=11, color=MUTED)

    document.add_paragraph()
    acceptance = document.add_paragraph(
        "Подтверждаю, что все 100 сценариев проверены в production Release, "
        "решения по строкам зафиксированы, а указанный итог соответствует фактической приёмке."
    )
    acceptance.paragraph_format.space_before = Pt(12)
    acceptance.paragraph_format.space_after = Pt(18)
    acceptance_run = acceptance.runs[0]
    set_run_font(acceptance_run, size=11, bold=True)

    for label in ("Владелец / ФИО", "Подпись", "Дата"):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after = Pt(8)
        label_run = paragraph.add_run(f"{label}: ")
        set_run_font(label_run, size=11, bold=True)
        line_run = paragraph.add_run("____________________________________________________________")
        set_run_font(line_run, size=11, color=MUTED)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    print(f"WROTE {output}")
    print(f"ROWS {len(rows)}")
    print("COUNTS " + " ".join(f"{key}={counts.get(key, 0)}" for key in ("PASS", "PARTIAL", "PENDING", "FAIL", "BLOCKED")))


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    build_document(args.source.resolve(), args.output.resolve())


if __name__ == "__main__":
    main()
