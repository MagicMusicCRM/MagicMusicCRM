from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_instruction():
    doc = Document()
    
    # Title
    title = doc.add_heading('Инструкция по настройке Custom Domain для Supabase', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Intro
    p = doc.add_paragraph()
    p.add_run('Данная настройка необходима для восстановления стабильной работы CRM системы MagicMusic на территории РФ. ').bold = True
    p.add_run('Использование выделенного поддомена школы позволяет обойти региональные блокировки сетевых адресов и восстановить работу чатов (WebSockets).')
    
    # Section 1
    doc.add_heading('1. Задача', level=1)
    doc.add_paragraph('Настроить поддомен (например, api.magic-music.org) так, чтобы он указывал на инфраструктуру Supabase. Это позволит приложению работать без VPN и сторонних утилит.')
    
    # Section 2
    doc.add_heading('2. Технические данные для настройки DNS', level=1)
    doc.add_paragraph('Ниже приведены ПРИМЕРНЫЕ данные. Актуальные значения (токены верификации) будут предоставлены после активации функции на стороне Supabase.')
    
    # Table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Тип записи'
    hdr_cells[1].text = 'Имя (Host)'
    hdr_cells[2].text = 'Значение (Value)'
    hdr_cells[3].text = 'Примечание'
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True

    # Row 1 (CNAME)
    row_cells = table.add_row().cells
    row_cells[0].text = 'CNAME'
    row_cells[1].text = 'api'
    row_cells[2].text = 'xbnywnlhfgofskbdxb.supabase.co'
    row_cells[3].text = 'Основной шлюз API'
    
    # Row 2 (TXT)
    row_cells = table.add_row().cells
    row_cells[0].text = 'TXT'
    row_cells[1].text = '_supabase.api'
    row_cells[2].text = 'supabase-verification=example-token-123'
    row_cells[3].text = 'Верификация владения'
    
    # Section 3
    doc.add_heading('3. Важное примечание (Cloudflare)', level=1)
    warn_p = doc.add_paragraph()
    warn_run = warn_p.add_run('Если домен magic-music.org обслуживается через Cloudflare, для данных записей необходимо выключить проксирование (режим "DNS Only", серое облако).')
    warn_run.bold = True
    
    # Section 4
    doc.add_heading('4. Порядок действий', level=1)
    steps = [
        'Администратор подтверждает возможность создания поддомена.',
        'Разработчик активирует функцию Custom Domain в Supabase.',
        'Разработчик передает актуальные CNAME и TXT токены.',
        'Администратор вносит записи в DNS.',
        'После верификации (от 15 мин до 2 часов) система начинает работать.'
    ]
    for step in steps:
        doc.add_paragraph(step, style='List Bullet')
        
    # Save
    filename = r'C:\Users\User\.gemini\antigravity\brain\dc49f24f-9f14-4129-b17c-1b1900f52498\artifacts\Supabase_Custom_Domain_Setup_Guide.docx'
    doc.save(filename)
    print(f"File saved to: {filename}")

if __name__ == "__main__":
    create_instruction()
